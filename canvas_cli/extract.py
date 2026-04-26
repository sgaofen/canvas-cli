"""Shared text extraction: PDF / DOCX / HTML → plaintext.

Used by both `canvas read` (single-file extract) and `canvas search --content`
(bulk index build). Pure functions on bytes/strings — no I/O.
"""

from __future__ import annotations

import io
import logging
import re

import html2text

# pypdf logs harmless "Ignoring wrong pointing object" warnings to stderr for
# many real-world PDFs. Mute it — extraction still succeeds.
logging.getLogger("pypdf").setLevel(logging.ERROR)


def extract_links(text: str) -> list[str]:
    """Pull all unique http(s) URLs out of plaintext / markdown."""
    if not text:
        return []
    raw = re.findall(r'https?://[^\s\)\]\>"\'<]+', text)
    seen = set()
    out: list[str] = []
    for u in raw:
        u = u.rstrip(".,;:!?")
        if u and u not in seen:
            seen.add(u)
            out.append(u)
    return out


def to_markdown(html: str) -> str:
    """Render HTML to Markdown via html2text (no line wrapping)."""
    if not html:
        return ""
    h = html2text.HTML2Text()
    h.body_width = 0
    h.ignore_images = False
    return h.handle(html).strip()


def extract_pdf(data: bytes) -> str:
    """Extract text from PDF bytes via pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[pypdf not installed]"
    try:
        reader = PdfReader(io.BytesIO(data))
        out: list[str] = []
        for i, page in enumerate(reader.pages, 1):
            out.append(f"\n--- page {i} ---\n")
            out.append(page.extract_text() or "")
        return "".join(out).strip()
    except Exception as exc:
        return f"[PDF extraction failed: {exc}]"


def extract_docx(data: bytes) -> str:
    """Extract text from DOCX bytes via python-docx (paragraphs + tables)."""
    try:
        from docx import Document
    except ImportError:
        return "[python-docx not installed]"
    try:
        doc = Document(io.BytesIO(data))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    except Exception as exc:
        return f"[DOCX extraction failed: {exc}]"


def extract_by_filename(filename: str, data: bytes) -> str:
    """Pick the right extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_pdf(data)
    if lower.endswith(".docx"):
        return extract_docx(data)
    if lower.endswith((".txt", ".md", ".csv", ".log", ".json", ".html", ".htm")):
        text = data.decode("utf-8", errors="replace")
        if lower.endswith((".html", ".htm")):
            return to_markdown(text)
        return text
    return ""
