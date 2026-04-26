"""`canvas read` — extract text from files, pages, assignments, or arbitrary URLs.

This is the core "let an AI agent read course materials" command. It routes
based on flags and returns plain text (or Markdown for HTML sources).

  canvas read --file <id>            # PDF / DOCX / TXT → text
  canvas read --page <slug> --course X  # course Page body → markdown
  canvas read --assignment <id> --from-course X
  canvas read --url <URL>            # arbitrary external (Google Docs, PDFs, ...)

Every JSON envelope includes a `links` array of `http(s)://` URLs found in
the extracted text — agents can use that to follow chains (Page → Google Doc
→ embedded PDF) without re-parsing themselves.
"""

from __future__ import annotations

import io
import json as _json
import re
from pathlib import Path

import html2text
import httpx
import typer
from canvasapi.exceptions import CanvasException, ResourceDoesNotExist
from rich.console import Console

from ..client import get_canvas, get_user_courses
from ..config import load_config
from ..extract import extract_links as _extract_links
from ..matchers import resolve_course

console = Console()
err_console = Console(stderr=True)


def _to_markdown(html: str) -> str:
    h = html2text.HTML2Text()
    h.body_width = 0
    return h.handle(html).strip() if html else ""


def _title_from_html(html: str, fallback: str = "?") -> str:
    m = re.search(r"<title[^>]*>([^<]+)</title>", html, re.IGNORECASE)
    if m:
        return m.group(1).strip()
    return fallback


def read(
    file_id: int | None = typer.Option(None, "--file", help="Canvas file id"),
    page_ref: str | None = typer.Option(None, "--page", help="Page url-slug or id"),
    course: str | None = typer.Option(
        None, "--course", "-c", help="Course (required when --page is used)"
    ),
    assignment_id: int | None = typer.Option(None, "--assignment", help="Assignment id"),
    course_for_assignment: str | None = typer.Option(
        None, "--from-course", help="Course (required when --assignment is used)"
    ),
    url: str | None = typer.Option(
        None, "--url", help="Fetch any external URL (Google Docs, PDF, HTML, ...)"
    ),
    json_out: bool = typer.Option(False, "--json", help="Emit JSON envelope"),
) -> None:
    """Extract text from a file, page, assignment, or arbitrary URL."""
    targets = sum(x is not None for x in (file_id, page_ref, assignment_id, url))
    if targets != 1:
        err_console.print(
            "[red]Specify exactly one of --file / --page / --assignment / --url[/red]"
        )
        raise typer.Exit(code=2)

    if file_id is not None:
        result = _read_file(file_id)
    elif url is not None:
        result = _read_url(url)
    elif page_ref is not None:
        if not course:
            err_console.print("[red]--page requires --course[/red]")
            raise typer.Exit(code=2)
        result = _read_page(course, page_ref)
    else:
        if not course_for_assignment:
            err_console.print("[red]--assignment requires --from-course[/red]")
            raise typer.Exit(code=2)
        result = _read_assignment(course_for_assignment, assignment_id)  # type: ignore[arg-type]

    # Always populate `links` so agents can chase chains
    result["links"] = _extract_links(result.get("text", ""))

    if json_out:
        print(_json.dumps(result))
    else:
        console.print(f"[bold cyan]{result.get('source')}[/bold cyan]: {result.get('title', '?')}")
        if result.get("url"):
            console.print(f"[dim]{result['url']}[/dim]")
        console.print()
        print(result.get("text", ""))
        if result["links"]:
            console.print(f"\n[dim]Found {len(result['links'])} link(s):[/dim]")
            for link in result["links"][:10]:
                console.print(f"  [dim]{link}[/dim]")
            if len(result["links"]) > 10:
                console.print(f"  [dim]... +{len(result['links']) - 10} more[/dim]")


def _read_file(file_id: int) -> dict:
    canvas = get_canvas()
    try:
        f = canvas.get_file(file_id)
    except ResourceDoesNotExist:
        err_console.print(f"[red]File not found:[/red] {file_id}")
        raise typer.Exit(code=1)
    except CanvasException as exc:
        err_console.print(f"[red]Failed:[/red] {exc}")
        raise typer.Exit(code=1)

    name = getattr(f, "display_name", "?") or "?"
    file_url = getattr(f, "url", "") or ""
    if not file_url:
        err_console.print("[red]No download URL on file[/red]")
        raise typer.Exit(code=1)

    local_path = _find_local_file(file_id)
    if local_path:
        data = local_path.read_bytes()
        source = f"local:{local_path}"
    else:
        with httpx.Client(timeout=httpx.Timeout(30.0, read=300.0), follow_redirects=True) as client:
            resp = client.get(file_url)
            resp.raise_for_status()
            data = resp.content
        source = f"canvas:file:{file_id}"

    text = _extract_by_filename(name, data)
    return {
        "source": source,
        "file_id": file_id,
        "title": name,
        "url": file_url,
        "size": len(data),
        "text": text,
        "extracted_chars": len(text),
    }


_GDOCS_EDIT_RE = re.compile(
    r"^(https://docs\.google\.com/(document|spreadsheets|presentation)/d/[^/]+)(/(edit|view|preview)?)?(\?.*)?$"
)


def _normalize_gdocs(url: str) -> str:
    """Translate Google Docs /edit URLs to /export endpoints for clean text fetch."""
    m = _GDOCS_EDIT_RE.match(url)
    if not m:
        return url
    base, kind = m.group(1), m.group(2)
    fmt = {"document": "txt", "spreadsheets": "csv", "presentation": "txt"}[kind]
    return f"{base}/export?format={fmt}"


def _read_url(target_url: str) -> dict:
    """Fetch any external URL and extract text based on content-type / extension."""
    fetch_url = _normalize_gdocs(target_url)
    try:
        with httpx.Client(
            timeout=httpx.Timeout(30.0, read=120.0),
            follow_redirects=True,
            headers={"User-Agent": "canvas-cli/0.1 (+https://github.com/local)"},
        ) as client:
            resp = client.get(fetch_url)
            resp.raise_for_status()
            data = resp.content
            ct = resp.headers.get("content-type", "").lower()
            final_url = str(resp.url)
    except httpx.HTTPStatusError as exc:
        return {
            "source": "external:url",
            "url": fetch_url,
            "requested_url": target_url,
            "title": target_url,
            "error": f"HTTP {exc.response.status_code}: {exc.response.reason_phrase}",
            "status_code": exc.response.status_code,
            "kind": "error",
            "size": 0,
            "text": "",
            "extracted_chars": 0,
        }
    except httpx.HTTPError as exc:
        return {
            "source": "external:url",
            "url": fetch_url,
            "requested_url": target_url,
            "title": target_url,
            "error": str(exc),
            "kind": "error",
            "size": 0,
            "text": "",
            "extracted_chars": 0,
        }

    lower_url = fetch_url.lower()
    title = target_url
    kind: str

    if "application/pdf" in ct or lower_url.endswith(".pdf"):
        text = _extract_pdf(data)
        kind = "pdf"
    elif "officedocument.wordprocessingml" in ct or lower_url.endswith(".docx"):
        text = _extract_docx(data)
        kind = "docx"
    elif "text/csv" in ct or "text/tab-separated-values" in ct or lower_url.endswith((".csv", ".tsv")):
        text = data.decode("utf-8", errors="replace")
        kind = "csv"
    elif "text/html" in ct or "html" in ct:
        try:
            html = data.decode(resp.encoding or "utf-8", errors="replace")
        except (LookupError, AttributeError):
            html = data.decode("utf-8", errors="replace")
        title = _title_from_html(html, fallback=target_url)
        text = _to_markdown(html)
        kind = "html"
    elif "text/" in ct or "json" in ct or "xml" in ct:
        text = data.decode("utf-8", errors="replace")
        kind = "text"
    else:
        text = f"[unsupported content-type: {ct}]"
        kind = "unknown"

    return {
        "source": "external:url",
        "url": final_url,
        "requested_url": target_url,
        "title": title,
        "content_type": ct,
        "kind": kind,
        "size": len(data),
        "text": text,
        "extracted_chars": len(text),
    }


def _find_local_file(file_id: int) -> Path | None:
    try:
        config = load_config()
    except SystemExit:
        return None
    sync_root = Path(config["sync_dir"]).expanduser()
    if not sync_root.exists():
        return None
    for manifest_path in sync_root.glob("*/.canvas-manifest.json"):
        try:
            data = _json.loads(manifest_path.read_text())
        except (OSError, ValueError):
            continue
        files_meta = data.get("files", {})
        meta = files_meta.get(str(file_id))
        if not meta:
            continue
        candidate = manifest_path.parent / meta.get("path", "")
        if candidate.exists():
            return candidate
    return None


def _extract_by_filename(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(data)
    if lower.endswith(".docx"):
        return _extract_docx(data)
    if lower.endswith((".txt", ".md", ".csv", ".log", ".json", ".html", ".htm")):
        try:
            text = data.decode("utf-8", errors="replace")
        except Exception:
            text = ""
        if lower.endswith((".html", ".htm")):
            return _to_markdown(text)
        return text
    return f"[unsupported file type: {filename}]"


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[pypdf not installed]"
    try:
        reader = PdfReader(io.BytesIO(data))
        out = []
        for i, page in enumerate(reader.pages, 1):
            out.append(f"\n--- page {i} ---\n")
            out.append(page.extract_text() or "")
        return "".join(out).strip()
    except Exception as exc:
        return f"[PDF extraction failed: {exc}]"


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[python-docx not installed]"
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    except Exception as exc:
        return f"[DOCX extraction failed: {exc}]"


def _read_page(course: str, page_ref: str) -> dict:
    courses = get_user_courses(active_only=True)
    target = resolve_course(courses, course)
    try:
        p = target.get_page(page_ref)
    except ResourceDoesNotExist:
        err_console.print(f"[red]Page not found:[/red] {page_ref}")
        raise typer.Exit(code=1)
    except CanvasException as exc:
        err_console.print(f"[red]Failed:[/red] {exc}")
        raise typer.Exit(code=1)

    html = getattr(p, "body", "") or ""
    md = _to_markdown(html)
    return {
        "source": "canvas:page",
        "page_id": getattr(p, "page_id", None),
        "title": getattr(p, "title", "?"),
        "url": getattr(p, "html_url", None) or getattr(p, "url", None),
        "text": md,
        "extracted_chars": len(md),
    }


def _read_assignment(course: str, assignment_id: int) -> dict:
    courses = get_user_courses(active_only=True)
    target = resolve_course(courses, course)
    try:
        a = target.get_assignment(assignment_id)
    except ResourceDoesNotExist:
        err_console.print(f"[red]Assignment not found:[/red] {assignment_id}")
        raise typer.Exit(code=1)
    except CanvasException as exc:
        err_console.print(f"[red]Failed:[/red] {exc}")
        raise typer.Exit(code=1)

    html = getattr(a, "description", "") or ""
    md = _to_markdown(html)
    return {
        "source": "canvas:assignment",
        "assignment_id": assignment_id,
        "title": getattr(a, "name", "?"),
        "url": getattr(a, "html_url", None),
        "due_at": getattr(a, "due_at", None),
        "points_possible": getattr(a, "points_possible", None),
        "text": md,
        "extracted_chars": len(md),
    }
