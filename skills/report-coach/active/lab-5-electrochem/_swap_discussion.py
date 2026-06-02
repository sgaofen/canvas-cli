"""Replace the 4 Discussion paragraphs (P109-P112) in Lab 5 Worksheet_FILLED.docx
with the user's polished version, preserving OMML math rendering via a pandoc-generated
fragment.

IL-2 compliant: surgical splice; we do NOT rebuild the whole docx from markdown.
"""
import sys, os, shutil, tempfile
sys.stdout.reconfigure(encoding='utf-8')
import pypandoc
from docx import Document
from copy import deepcopy
from lxml import etree

DOCX = r"D:\Desktop\m3lc\week 5\Lab 5 Worksheet_FILLED.docx"
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
def qn(tag): return f'{{{W_NS}}}{tag}'

# Stephen's polished Discussion text — transcribed exactly with LaTeX math.
POLISHED_MD = r"""The voltage measured by the cell is then equal to $E_\mathrm{Ag} = E^\circ_{\mathrm{Ag^+/Ag}} + 0.05916 \log[\mathrm{Ag^+}]$ and depends on the free $[\mathrm{Ag^+}]$ in solution, which is related to the free $[\mathrm{Cl^-}]$ by the equilibrium $[\mathrm{Ag^+}][\mathrm{Cl^-}] = K_\text{sp}$.

Prior to the equivalence point, essentially all of the $\mathrm{Ag^+}$ from the burette precipitates as AgCl with the excess of $\mathrm{Cl^-}$ present. The concentration of $[\mathrm{Ag^+}]$ remains very small (established by $K_\text{sp}/[\mathrm{Cl^-}]$) and $E_\text{cell}$ increases gradually as $[\mathrm{Cl^-}]$ decreases.

At the equivalence point, $[\mathrm{Ag^+}] = [\mathrm{Cl^-}] = \sqrt{K_\text{sp}}$. An extra small volume of titrant will react with the remaining chloride, and the voltage will increase rapidly by orders of magnitude over a small volume at $V_\text{eq}$ — just a few mL.

Once the equivalence point is reached, all of the chloride has been consumed and the excess $\mathrm{Ag^+}$ remains in solution as the predominant species. Now the $[\mathrm{Ag^+}]$ increases approximately in a linear fashion, and $E_\text{cell}$ continues to increase, albeit at the Nernstian rate of $\sim\!59$ mV per decade of $[\mathrm{Ag^+}]$.
"""

# 1. Run pandoc on the snippet to a temp docx.
tmp_dir = tempfile.mkdtemp()
tmp_md = os.path.join(tmp_dir, "polished.md")
tmp_docx = os.path.join(tmp_dir, "polished.docx")
with open(tmp_md, 'w', encoding='utf-8') as f:
    f.write(POLISHED_MD)
pypandoc.convert_file(tmp_md, 'docx', outputfile=tmp_docx,
    extra_args=["--from=markdown+tex_math_dollars+tex_math_single_backslash", "--to=docx", "--standalone"])
print(f"[1] pandoc -> {tmp_docx}")

# 2. Read the pandoc output docx; grab its 4 body paragraphs (skip empty ones).
src_doc = Document(tmp_docx)
new_pars = []
for p in src_doc.paragraphs:
    if p.text.strip():
        new_pars.append(p._element)
        # peek
        print(f"  new paragraph: {p.text[:60]!r}...")
print(f"[2] {len(new_pars)} paragraphs extracted from pandoc output")
assert len(new_pars) == 4, f"Expected 4 paragraphs, got {len(new_pars)}"

# 3. Open the destination docx, locate the 4 Discussion paragraphs to replace.
dst_doc = Document(DOCX)
dst_paras = list(dst_doc.paragraphs)
# Find P109-P112 by indices; verify content first.
target_idxs = []
for i, p in enumerate(dst_paras):
    t = p.text.strip()
    if t.startswith("The silver indicator electrode reads") or \
       t.startswith("Before the equivalence point,") or \
       t.startswith("At the equivalence point,") or \
       t.startswith("After the equivalence point,"):
        target_idxs.append(i)
        print(f"  target P{i}: {t[:60]!r}")
print(f"[3] Found {len(target_idxs)} targets")
assert len(target_idxs) == 4, "Could not locate all 4 Discussion paragraphs"

# 4. Splice: insert new paragraphs before the first target, then delete the originals.
first_target = dst_paras[target_idxs[0]]._element
parent = first_target.getparent()

for new_p in new_pars:
    # Deep copy so we don't mess with the source tree
    parent.insert(list(parent).index(first_target), deepcopy(new_p))

# Now remove the originals.
for idx in target_idxs:
    p = dst_paras[idx]._element
    p.getparent().remove(p)

print("[4] Splice complete; saving...")
dst_doc.save(DOCX)
print(f"Saved: {DOCX}")
shutil.rmtree(tmp_dir, ignore_errors=True)
