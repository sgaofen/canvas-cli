"""
Fill Lab 3 worksheet answers + figures + captions + Methods section.

Design:
- Preserve all original question text and paragraph styling.
- Insert new paragraphs after each question-identifying fingerprint.
- Figures inserted as images; captions inserted as SEPARATE paragraphs below (per
  TA rule 2026-04-18: "You will lose credit if your figure has a title built into it").
- Methods section inserted before "Data Analysis" heading.
- Output to _FILLED.docx so original is preserved.
"""
import os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

BASE = r"C:\Users\YY\.claude\skills\report-coach\active\lab-3-riboflavin"
SRC = os.path.join(BASE, "Lab 3 Worksheet_ Determination of Riboflavin Content in A Multivitamin.docx")
OUT = os.path.join(BASE, "Lab 3 Worksheet_FILLED.docx")

doc = Document(SRC)
ANSWER_BLUE = RGBColor(0x1F, 0x49, 0x7D)


def insert_paragraph_after(paragraph, text="", italic=False, bold=False, color=None, style=None):
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = doc.styles[style]
        except KeyError:
            pass
    if text:
        run = new_para.add_run(text)
        if italic: run.italic = True
        if bold: run.bold = True
        if color: run.font.color.rgb = color
    return new_para


def insert_image_after(paragraph, image_path, width_in=6.0):
    new_para = insert_paragraph_after(paragraph)
    new_para.add_run().add_picture(image_path, width=Inches(width_in))
    return new_para


def append_block(after_para, items):
    last = after_para
    for kind, payload, kwargs in items:
        if kind == "text":
            last = insert_paragraph_after(
                last, payload,
                italic=kwargs.get("italic", False),
                bold=kwargs.get("bold", False),
                color=kwargs.get("color"),
                style=kwargs.get("style"),
            )
        elif kind == "image":
            last = insert_image_after(last, payload, width_in=kwargs.get("width_in", 6.0))
        elif kind == "blank":
            last = insert_paragraph_after(last, "")
    return last


# ---------------------------------------------------------------------------
# Methods section — inserted BEFORE "Data Analysis" heading
# ---------------------------------------------------------------------------
METHODS = [
    ("text", "Methods", {"bold": True, "style": "Heading 1"}),
    ("text",
     "Riboflavin content of a powdered multivitamin supplement was quantified by 405 nm-excitation fluorescence "
     "spectroscopy on a Vernier SpectroVis Plus spectrophotometer. All solutions were prepared in 0.02 M aqueous "
     "acetic acid, which was also used as the blank and as the diluent throughout. A 50.0 ppm riboflavin stock "
     "was dissolved in the acid and diluted into a seven-point calibration series spanning 2–20 ppm (25 mL "
     "volumetric flasks); a separate multivitamin sample was powdered, weighed, and dissolved in the same "
     "solvent to produce the unknown solution.",
     {}),
    ("text",
     "Absorbance spectra of the calibration series and the unknown were collected first (380–700 nm) and used to "
     "select the excitation wavelength by comparing the absorbance of the most concentrated standard at 405 nm "
     "and at 500 nm. Fluorescence emission spectra (405 nm excitation) were then recorded for each standard, for "
     "three independent aliquots of the undiluted unknown, and for a five-point standard-addition series in which "
     "5.0 mL of the unknown was combined with 0, 1, 2, 3, or 4 mL of the 50.0 ppm stock and diluted to 25.0 mL. "
     "Peak emission intensity near 517 nm served as the analytical signal. Between samples the cuvette was rinsed "
     "with the next solution and oriented consistently relative to the beam.",
     {}),
    ("text",
     "Riboflavin concentration in the unknown was determined by two independent methods: (i) an external "
     "calibration curve of peak intensity vs. concentration, restricted to the 2–12 ppm linear region (the 20 ppm "
     "standard was excluded for inner-filter-effect plateauing), and (ii) the method of standard additions, with "
     "the analyte concentration recovered from the x-intercept of peak intensity vs. added stock volume. "
     "Ninety-five-percent confidence intervals were propagated using the standard single-analyte calibration "
     "formulas in Miller & Miller.",
     {}),
    ("blank", "", {}),
]

# ---------------------------------------------------------------------------
# Per-question answer payloads
# ---------------------------------------------------------------------------

ANSWERS = [
    # Q1 - absorbance overlay
    ("overlaying the collected absorbance spectra of the standards", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("image", os.path.join(BASE, "fig1_absorbance_overlay.png"), {"width_in": 6.2}),
        ("text",
         "Figure 1. Absorbance spectra of seven riboflavin calibration standards (2, 4, 6, 8, 10, 12, and 20 ppm "
         "in 0.02 M acetic acid) and the unknown multivitamin solution (red dashed trace). All spectra show the "
         "characteristic riboflavin absorption maximum at λmax = 440.6 nm (grey dotted vertical line). The "
         "unknown lies between the 6 ppm and 10 ppm standards and retains the same peak position and shape as "
         "pure riboflavin, confirming that riboflavin is the dominant visible-range absorber in the tablet.",
         {"italic": True}),
        ("blank", "", {}),
    ]),

    # Q2 - excitation wavelength
    ("which wavelength (405 nm or 500 nm) should be used as the excitation", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "405 nm should be used as the excitation wavelength.",
         {"bold": True}),
        ("text",
         "At the 20 ppm standard the absorbance at 405 nm is 0.413 compared to 0.058 at 500 nm — 405 nm absorbs "
         "about seven times more strongly. A stronger excitation absorbance produces a stronger fluorescence "
         "signal, so 405 nm gives a larger dynamic range for the calibration curve and better signal-to-noise at "
         "the low concentrations of interest. The absorption maximum is at 440.6 nm, but the spectrophotometer "
         "only offers 405 nm and 500 nm as excitation options, and 405 nm is on the steep rising edge of the "
         "riboflavin absorption band while 500 nm is well past the peak into the tail.",
         {}),
        ("blank", "", {}),
    ]),

    # Q3 - fluorescence overlay standards + unknown trials
    ("overlaying the spectra of all the standards and the unknown trials", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("image", os.path.join(BASE, "fig2_fluorescence_overlay.png"), {"width_in": 6.2}),
        ("text",
         "Figure 2. Fluorescence emission spectra (405 nm excitation) of the seven riboflavin calibration "
         "standards (2–20 ppm) and three independent aliquots of the undiluted multivitamin unknown (red, orange, "
         "and brown dashed traces, respectively). All spectra share the characteristic riboflavin emission "
         "maximum near λmax ≈ 517 nm. The three unknown trials overlap almost exactly, indicating high "
         "measurement precision; their peak intensities sit between those of the 6 ppm and 10 ppm standards.",
         {"italic": True}),
        ("blank", "", {}),
    ]),

    # Q4a - calibration all 7 points
    ("calibration curve (fluorescence intensity vs. concentration) for all the standards", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("image", os.path.join(BASE, "fig3_cal_all7.png"), {"width_in": 5.8}),
        ("text",
         "Figure 3. Fluorescence calibration curve of peak intensity at ~517 nm vs. riboflavin concentration for "
         "all seven standards. The linear fit is y = 0.01008x + 0.01027 with R² = 0.985. The reduced R² reflects "
         "deviation of the highest (20 ppm) standard, which falls ~13% below the line defined by the lower "
         "concentrations — a signature of the inner-filter effect.",
         {"italic": True}),
        ("text",
         "The most linear concentration range is 2–12 ppm. Over this range the residuals of a linear fit are "
         "small (≤ 0.006 intensity units, or ≤ 5% of signal) and R² = 0.994; once 20 ppm is included the fit "
         "flattens because that point is held back by inner-filter attenuation and no longer scales linearly "
         "with concentration.",
         {}),
        ("blank", "", {}),
    ]),

    # Q4b - exclude inner filter
    ("determine which standard solutions are too concentrated and regraph", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Yes — the 20 ppm standard shows the inner-filter effect.",
         {"bold": True}),
        ("text",
         "Evidence: fitting only the 2–12 ppm points gives the line y = 0.01168x + 0.00050 (R² = 0.994). Using "
         "that line to predict the signal at 20 ppm gives 0.234, but the measured intensity is 0.203 — about 13% "
         "lower than extrapolation, and the single largest residual by a factor of more than four. In a sample "
         "that concentrated, the excitation beam is absorbed strongly in the front ~1 mm of the cuvette and "
         "cannot reach the bulk, so the volume that fluoresces shrinks and the signal saturates. The regraphed "
         "calibration below excludes the 20 ppm point.",
         {}),
        ("image", os.path.join(BASE, "fig4_cal_linear.png"), {"width_in": 5.8}),
        ("text",
         "Figure 4. Fluorescence calibration curve after excluding the 20 ppm standard (shown as an open red "
         "circle) that exhibits the inner-filter effect. Linear fit over the retained 2–12 ppm range: "
         "y = 0.01168x + 0.00050 with R² = 0.994. This line is used for all unknown-concentration calculations.",
         {"italic": True}),
        ("blank", "", {}),
    ]),

    # Q5a - slope / intercept with 95% CI
    ("Report the slope and y-intercept of the more linear calibration curve", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "From the 2–12 ppm linear fit (n = 6 standards, df = 4, t0.05,4 = 2.776):",
         {}),
        ("text",
         "Unrounded statistics — slope m = 0.01168 ppm⁻¹, 95% CI = ±0.00129; intercept b = 0.00050, "
         "95% CI = ±0.01005. Fluorescence intensity is dimensionless (no \"AU\"), so the intercept has no unit.",
         {}),
        ("text",
         "Applying the sig-fig rule (error rounded to 1 sig fig at the first non-zero digit; value rounded to "
         "the same decimal place):",
         {}),
        ("text", "Slope:     m = 0.012 ± 0.001 ppm⁻¹", {"bold": True}),
        ("text", "Intercept: b = 0.00 ± 0.01",        {"bold": True}),
        ("blank", "", {}),
    ]),

    # Q5b - unknown concentration (cal) with 95% CI
    ("Calculate and report the undiluted unknown riboflavin concentration (in the vitamin solution made in Step A.3) using the more linear", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Three replicate measurements of the undiluted unknown gave peak intensities 0.0840, 0.0847, and "
         "0.0851; mean Ī = 0.08460, sample standard deviation s = 5.4×10⁻⁴. The unknown was not diluted before "
         "the fluorescence measurement, so the concentration in the cuvette equals the concentration in the "
         "Part A.3 solution.",
         {}),
        ("text",
         "c_A.3 = (Ī − b) / m  =  (0.08460 − 0.00050) / 0.01168  =  7.2023 ppm  (carried unrounded into the CI below)",
         {}),
        ("text",
         "95% CI by the single-analyte calibration formula (Miller & Miller), with s_yx = 3.89×10⁻³, "
         "Σ(xi−x̄)² = 70.0 ppm², N = 3 unknown replicates, n = 6 standards, t0.05,4 = 2.776:",
         {}),
        ("text",
         "s_c = (s_yx / m) · √(1/N + 1/n + (Ī − ȳ_cal)² / (m² · Σ(xi−x̄)²))  =  0.2356 ppm",
         {}),
        ("text",
         "95% CI = t · s_c = 2.776 × 0.2356 = ±0.654 ppm",
         {}),
        ("text",
         "Rounding the CI to 1 sig fig (0.7) and the value to the same decimal place:",
         {}),
        ("text", "[Riboflavin] in the A.3 unknown solution = 7.2 ± 0.7 ppm", {"bold": True}),
        ("blank", "", {}),
    ]),

    # Q5c - mass percent (cal)
    ("Calculate the mass percent of riboflavin in the vitamin sample determined from the calibration curve", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Mass of riboflavin in the 250.0 mL A.3 solution, using the unrounded concentration:",
         {}),
        ("text",
         "m_riboflavin = c × V = 7.2023 mg/L × 0.2500 L = 1.8006 mg",
         {}),
        ("text",
         "Part A.3 was prepared by our partner group (each pair is assigned either step A.2 stock or step A.3 "
         "unknown, and the solutions are shared). Their notebook-recorded multivitamin-powder mass was not "
         "available at the time of writing; the manual's target of 50.0 mg is used below as the best available "
         "estimate and can be swapped in once the partner group's exact mass is obtained.",
         {"italic": True}),
        ("text",
         "mass% = m_riboflavin / m_vitamin × 100% = 1.8006 mg / 50.0 mg × 100% = 3.601%",
         {}),
        ("text", "Mass percent of riboflavin (calibration method) ≈ 3.6%", {"bold": True}),
        ("blank", "", {}),
    ]),

    # Q6 - SA fluorescence overlay
    ("overlaying the fluorescence spectra of all the standard addition solutions", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("image", os.path.join(BASE, "fig5_sa_spectra.png"), {"width_in": 6.2}),
        ("text",
         "Figure 5. Fluorescence emission spectra (405 nm excitation) of the five standard-addition solutions. "
         "Each solution contained 5.0 mL of the A.3 multivitamin unknown plus the indicated volume (0–4 mL) of "
         "50.0 ppm riboflavin stock, diluted to 25.0 mL with 0.02 M acetic acid. The emission maximum stays near "
         "λmax ≈ 517 nm across all additions, and the peak intensity increases monotonically with added stock "
         "volume except for SA-3 (3 mL, see Q7a).",
         {"italic": True}),
        ("blank", "", {}),
    ]),

    # Q7a - SA linear curve
    ("Graph the standard addition linear curve of fluorescence intensity vs. added volume", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("image", os.path.join(BASE, "fig6_sa_linear.png"), {"width_in": 5.8}),
        ("text",
         "Figure 6. Standard-addition plot of peak fluorescence intensity (405 nm excitation) vs. mL of 50.0 ppm "
         "riboflavin stock added. The five-point linear fit gives y = 0.01925x + 0.02333 with R² = 0.942. The "
         "SA-3 point (3 mL addition) sits approximately 13% below the line and has a residual about five times "
         "larger than the other four points, suggesting an error in preparing that single flask. Excluding SA-3, "
         "the four-point fit is y = 0.02094x + 0.02333 with R² = 0.993; this cleaner line and its x-intercept "
         "(−1.114 mL, purple ×) are used for the concentration calculation in Q8.",
         {"italic": True}),
        ("blank", "", {}),
    ]),

    # Q7b - SA slope / intercept with 95% CI
    ("Report the slope and y-intercept of the standard addition linear curve", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Using the four-point SA fit excluding the SA-3 outlier (n = 4, df = 2, t0.05,2 = 4.303):",
         {}),
        ("text",
         "Unrounded statistics — slope m = 0.02094 mL⁻¹, 95% CI = ±0.00537; intercept b = 0.02333, "
         "95% CI = ±0.01231.",
         {}),
        ("text",
         "For completeness, the full five-point fit is slope = 0.01925 ± 0.00874 mL⁻¹ and intercept = "
         "0.02333 ± 0.02142; SA-3 clearly inflates the CIs. Rounding per the sig-fig rule (error to 1 sig fig, "
         "value to the same decimal place):",
         {}),
        ("text", "Slope:     m = 0.021 ± 0.005 mL⁻¹", {"bold": True}),
        ("text", "Intercept: b = 0.02 ± 0.01",        {"bold": True}),
        ("blank", "", {}),
    ]),

    # Q8a - unknown concentration (SA) with 95% CI
    ("Calculate and report the undiluted unknown riboflavin concentrations (in the vitamin solution made in Step A.3) from the standard addition", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "For the standard addition, each cuvette solution contains 5.0 mL of A.3 unknown diluted into 25.0 mL "
         "total (a 5× dilution) plus added 50.0 ppm stock. Setting total signal to zero gives an x-intercept "
         "V_x, from which the A.3 concentration is recovered as",
         {}),
        ("text",
         "c_A.3 = c_stock × |V_x| / V_unknown = 50.0 ppm × |V_x| / 5.0 mL = 10 × |V_x| ppm/mL",
         {}),
        ("text",
         "Using the unrounded intercept V_x = −b/m = −0.02333 / 0.02094 = −1.1143 mL:",
         {}),
        ("text",
         "c_A.3 = 10 × 1.1143 = 11.14 ppm",
         {}),
        ("text",
         "95% CI on V_x via the Miller–Miller x-intercept formula. For the four retained additions "
         "(x = 0, 1, 2, 4 mL) the relevant regression statistics are x̄ = 1.75 mL, "
         "Σ(xi−x̄)² = 8.75 mL², ȳ = 0.0600, s_yx = 3.70×10⁻³, with t0.05,2 = 4.303:",
         {}),
        ("text",
         "s_Vx = (s_yx / m) · √(1/n + ȳ² / (m² · Σ(xi−x̄)²)) = 0.193 mL  →  CI_Vx = ±0.83 mL",
         {}),
        ("text",
         "Propagating the ×10 factor (volumes and stock concentration treated as exact): "
         "CI on c_A.3 = ±8.3 ppm. Rounding error to 1 sig fig and value to matching decimal place:",
         {}),
        ("text", "[Riboflavin] in the A.3 unknown solution (standard addition) = 11 ± 8 ppm", {"bold": True}),
        ("blank", "", {}),
    ]),

    # Q8b - mass percent (SA)
    ("Calculate the mass percent of riboflavin in the vitamin sample determined from the standard addition", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Using the unrounded SA concentration:",
         {}),
        ("text",
         "m_riboflavin = 11.14 mg/L × 0.2500 L = 2.786 mg",
         {}),
        ("text",
         "Using the same m_vitamin estimate as Q5c (50.0 mg, manual target; the partner group's exact "
         "notebook value will be substituted when available):",
         {"italic": True}),
        ("text",
         "mass% = 2.786 mg / 50.0 mg × 100% = 5.571%",
         {}),
        ("text", "Mass percent of riboflavin (standard-addition method) ≈ 5.6%", {"bold": True}),
        ("blank", "", {}),
    ]),

    # Q9 - theoretical mass% and % error (with TA announcement caveat)
    ("multivitamin table information provided in class to determine the theoretical mass percent", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "As of submission, the nutrition-label information for the assigned multivitamin had not yet been "
         "released to students. Per the Canvas announcement of 2026-04-18, no penalty applies if that information "
         "is unavailable. The framework below can be populated with the label data once provided:",
         {"italic": True}),
        ("text",
         "theoretical mass% = (label mg of riboflavin per tablet) / (mass of one tablet in mg) × 100%",
         {}),
        ("text",
         "% error (calibration)       = |3.60% − theoretical%| / theoretical% × 100%",
         {}),
        ("text",
         "% error (standard addition) = |5.57% − theoretical%| / theoretical% × 100%",
         {}),
        ("text",
         "Illustrative calculation assuming a typical high-potency B-complex formulation "
         "(25 mg riboflavin per 500 mg tablet → theoretical mass% = 5.00%):",
         {"italic": True}),
        ("text",
         "% error (calibration)       = |3.60 − 5.00| / 5.00 × 100% ≈ 28%",
         {}),
        ("text",
         "% error (standard addition) = |5.57 − 5.00| / 5.00 × 100% ≈ 11%",
         {}),
        ("text",
         "The final numbers will be recalculated once the actual nutrition label is released.",
         {"italic": True}),
        ("blank", "", {}),
    ]),

    # Q10 - which method is theoretically better
    ("Theoretically, which method is better (calibration curve or standard addition curve)", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text", "Standard addition is the theoretically better method here.", {"bold": True}),
        ("text",
         "An external calibration curve relies on a critical assumption: the fluorescence response per ppm of "
         "analyte in the pure standards is the same as the response per ppm in the unknown. This assumption "
         "fails when the unknown contains a chemically complex matrix — as a multivitamin does, with other "
         "vitamins, minerals (Fe, Cu, Zn), binders, dyes, and coatings. These components can absorb some of the "
         "excitation light (primary inner-filter effect), re-absorb emitted fluorescence (secondary inner-filter "
         "effect), or quench the riboflavin excited state through collisional or paramagnetic mechanisms. "
         "Calibration in pure acetic-acid solvent cannot see any of those effects, so it systematically "
         "mis-estimates the concentration.",
         {}),
        ("text",
         "Standard addition performs the calibration in the sample's own matrix. Every SA flask contains the "
         "same 5.0 mL of unknown matrix, so whatever matrix-induced change in fluorescence efficiency exists, it "
         "is present equally across all five points. The slope of intensity vs. added stock therefore reflects "
         "the instrument response under the actual measurement conditions, and the x-intercept recovers the "
         "original unknown concentration with matrix effects built in. The trade-off is statistical precision: "
         "because SA is a single-unknown method and the fit is short-range, the 95% CI on the result is wider "
         "than for a well-populated external calibration curve — reflected in this experiment by 7.2 ± 0.7 ppm "
         "from calibration vs. 11 ± 8 ppm from standard addition.",
         {}),
        ("text",
         "Our data are consistent with a matrix effect in the direction expected for a multivitamin: the SA "
         "concentration (~11 ppm) is higher than the calibration concentration (~7.2 ppm), which implies that "
         "the matrix reduces fluorescence efficiency so that the calibration line in pure solvent "
         "under-predicts the true amount of riboflavin in the tablet.",
         {}),
        ("blank", "", {}),
    ]),

    # Q11 - Why is fluorescence more sensitive than absorbance
    ("Why is fluorescence spectroscopy more sensitive than absorbance spectroscopy", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Absorbance is a small decrease measured against a bright background. The detector always sees the "
         "full source intensity I₀; adding analyte lowers the transmitted intensity to I, and the absorbance "
         "A = log(I₀/I) encodes a small difference between two large numbers. Shot noise and source-fluctuation "
         "noise scale with the magnitude of I₀, so at low concentrations the signal change is buried in baseline "
         "noise.",
         {}),
        ("text",
         "Fluorescence is a small increase measured against zero background. With the detector oriented "
         "perpendicular to the excitation beam, no direct source light reaches it; in the absence of a "
         "fluorophore the reading is essentially zero. A single emitted photon adds to near-zero baseline, and "
         "modern photomultipliers/PMTs can detect that signal easily. The manual's candle analogy is apt: "
         "noticing one lit candle in a dark room is easy, but noticing one unlit candle in a room of 100 lit "
         "ones is very hard.",
         {}),
        ("text",
         "Quantitatively, fluorescence can routinely reach picomolar or lower detection limits, while absorbance "
         "bottoms out near A ≈ 0.001 (millimolar-ish for most dyes) before shot noise dominates — a difference "
         "of several orders of magnitude.",
         {}),
        ("blank", "", {}),
    ]),

    # Q12 - Why can't we use absorbance for multivitamin riboflavin?
    ("Why can", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Absorbance is not selective. Riboflavin absorbs in the visible near 440 nm, but so do many other "
         "components of a multivitamin tablet: β-carotene and other carotenoids (vitamin A precursors), "
         "iron-containing pigments, yellow and orange food dyes, and any colored binder all have absorption "
         "bands that overlap the riboflavin peak. The absorbance measured at 441 nm is therefore the sum of "
         "riboflavin's contribution and everything else's — you cannot deconvolve them from a single spectrum.",
         {}),
        ("text",
         "This is visible in the data: the unknown's apparent riboflavin concentration from its 441 nm "
         "absorbance is ~9 ppm, but the fluorescence-derived concentration is only ~7 ppm (calibration) or "
         "~11 ppm (SA); the absorbance estimate sits between these because it includes non-riboflavin absorbers. "
         "Fluorescence is selective because only molecules with a non-negligible fluorescence quantum yield "
         "emit, and most multivitamin components either don't fluoresce at all or don't fluoresce in "
         "riboflavin's 517 nm window — so riboflavin's emission signal is essentially uncontaminated by the "
         "matrix.",
         {}),
        ("blank", "", {}),
    ]),

    # Q13 - λex vs λem
    ("Which is larger, the excitation wavelength", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "The emission wavelength (λem) is larger than the excitation wavelength (λex).",
         {"bold": True}),
        ("text",
         "When a molecule absorbs an excitation photon at λex, it is promoted to a vibrationally excited level "
         "of the first excited singlet state. Before emitting, it relaxes non-radiatively to the lowest "
         "vibrational level of that excited state (internal conversion + vibrational relaxation), losing a small "
         "amount of energy as heat. The emitted photon therefore has lower energy than the absorbed photon. "
         "Because E = hc/λ, lower photon energy means longer wavelength, so λem > λex. The wavelength difference "
         "is the Stokes shift. For riboflavin in this experiment, λex = 405 nm and λem ≈ 517 nm — a Stokes shift "
         "of about 112 nm.",
         {}),
        ("blank", "", {}),
    ]),

    # Q14 - detector position
    ("Which position (a,b,c) would it be acceptable", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Either position a (above the sample) or position c (below the sample) is acceptable; position b is "
         "not.",
         {"bold": True}),
        ("text",
         "Positions a and c are both oriented at 90° to the excitation beam, so the detector never looks "
         "directly into the source. In that geometry the baseline signal is near zero, and the fluorescence "
         "photons that escape the sample sideways are the only thing reaching the detector — exactly the "
         "situation that makes fluorescence so sensitive. Standard benchtop fluorometers use this right-angle "
         "arrangement for that reason.",
         {}),
        ("text",
         "Position b is directly in line with the transmitted excitation beam. A detector placed there is "
         "dominated by the source's own photons (many orders of magnitude brighter than the fluorescence), and "
         "any emission signal would be buried under the transmitted light. Position b is appropriate for "
         "measuring absorbance, not fluorescence.",
         {}),
        ("blank", "", {}),
    ]),

    # Q15 - Absorbance vs fluorescence
    ("Using the following diagrams and in your own words, explain the difference", [
        ("text", "Answer:", {"italic": True, "color": ANSWER_BLUE}),
        ("text",
         "Absorbance: a photon with the right energy promotes an electron from the HOMO to the LUMO (upward "
         "arrow on the left diagram). The measurement is made at that single wavelength, and what we record is "
         "the decrease in transmitted source light. Whatever the molecule does with the absorbed energy — "
         "dissipate it as heat, undergo a chemical reaction, emit a photon — does not affect the absorbance; "
         "every molecule that absorbs contributes equally.",
         {}),
        ("text",
         "Fluorescence: the same HOMO → LUMO absorption happens first (the upward arrow on the right diagram). "
         "Before relaxing back, the molecule loses a small amount of energy within the LUMO manifold through "
         "vibrational relaxation, so the downward arrow (emission) is shorter than the upward arrow. The "
         "detector measures this newly emitted photon — at a wavelength different from the excitation beam, "
         "against an essentially zero background.",
         {}),
        ("text",
         "Three practical consequences follow:",
         {}),
        ("text",
         "1) Selectivity — every absorbing molecule contributes to absorbance, but only molecules with a "
         "non-negligible fluorescence quantum yield contribute to fluorescence. Fluorescence therefore isolates "
         "a much smaller subset of species.",
         {}),
        ("text",
         "2) Sensitivity — absorbance measures a small difference (I₀ − I) between two large numbers, while "
         "fluorescence measures a small signal on near-zero background. The latter is fundamentally easier, "
         "which is why fluorescence reaches single-molecule detection limits in modern instruments.",
         {}),
        ("text",
         "3) Wavelength — absorbance uses one wavelength (the detector sees the source), fluorescence uses two "
         "(λex ≠ λem, separated by the Stokes shift), allowing spectral discrimination of the signal from the "
         "excitation beam.",
         {}),
        ("blank", "", {}),
    ]),
]


# ---------------------------------------------------------------------------
# Phase 1: insert Methods section BEFORE the "Data Analysis" heading
# ---------------------------------------------------------------------------
for p in doc.paragraphs:
    if p.text.strip() == "Data Analysis":
        # Insert the methods block before this paragraph by locating the predecessor
        # Approach: walk backwards from Data Analysis and insert after the immediately
        # preceding non-empty paragraph.
        predecessor = None
        for q in doc.paragraphs:
            if q._element is p._element:
                break
            predecessor = q
        if predecessor is None:
            predecessor = p  # fallback: insert after anyway
        append_block(predecessor, METHODS)
        break


# ---------------------------------------------------------------------------
# Phase 2: fill answers for each question
# ---------------------------------------------------------------------------
# Must re-iterate because we inserted paragraphs. Use text-based lookup each time.
for fingerprint, items in ANSWERS:
    target = None
    for p in doc.paragraphs:
        if fingerprint in p.text:
            target = p
            break
    if target is None:
        print(f"[WARN] fingerprint not found: {fingerprint!r}")
        continue
    append_block(target, items)

doc.save(OUT)
print(f"[OK] wrote {OUT}")
